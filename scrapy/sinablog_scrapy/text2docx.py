# -*- coding: utf-8 -*-  
from docx import Document
from docx.shared import Inches
from sinablog_scrapy import settings
import os
from os.path import isfile, join
from docx.enum.text import WD_ALIGN_PARAGRAPH

def text2docx():
    doc = Document()
    filename = ""
    with open( os.path.join( settings.TXT_STORE, '100001.meta' ), 'r' ) as f:
        filename = f.readline().strip()
        if filename == None or len( filename ) == 0:
            print "Failed to find the blog meta data"
        title = f.readline().decode('UTF-8')
        content = f.readlines()
        content = u''.join( [ i.decode('UTF-8') for i in content if len( i.decode('UTF-8').strip() ) > 0  ] )
        p = doc.add_heading( title, 2 )
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph( content )

    files = [f for f in os.listdir(settings.TXT_STORE) if ( 'img' in f or 'text' in f ) and isfile(join(settings.TXT_STORE, f))]
    sorted( files )

    is_prev_img = False
    for afile in files:
        if 'img' in afile:
            img_file = open( os.path.join( settings.TXT_STORE, afile ), 'r' ).readline().strip()
            doc.add_picture( os.path.join( settings.IMAGES_STORE, img_file  ) )
            is_prev_img = True
        elif 'text' in afile:
            with open( os.path.join( settings.TXT_STORE, afile ), 'r' ) as f:
                print afile
                content = f.readlines()
                content = u''.join( [ i.decode('UTF-8') for i in content if len( i.decode('UTF-8').strip() ) > 0  ] )
                p =  doc.add_paragraph( content.strip() )
                paragraph_format = p.paragraph_format
                paragraph_format.first_line_indent = Inches(0.25)
                if is_prev_img:
                    is_prev_img = False
                    if len( content ) < 36:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        paragraph_format.first_line_indent = None

    doc.save( os.path.join( settings.DOCX_STORE, filename ))

text2docx()


